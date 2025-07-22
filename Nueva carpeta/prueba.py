import os
from docx import Document
import requests
import re
import json

def normalizar_texto(texto):
    return re.sub(r'\W+', '', texto.lower())

def extraer_lineas_y_tablas(doc):
    lineas = []
    for p in doc.paragraphs:
        if p.text.strip():
            lineas.append(p.text.strip())

    tabla_valores = {}
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = fila.cells
            for idx, celda in enumerate(celdas):
                etiqueta = celda.text.strip()
                if not etiqueta:
                    continue
                etiqueta_norm = normalizar_texto(etiqueta)

                valor = ""
                if idx + 1 < len(celdas):
                    valor = celdas[idx + 1].text.strip()
                if not valor and idx - 1 >= 0:
                    valor = celdas[idx - 1].text.strip()

                if valor:
                    tabla_valores[etiqueta_norm] = valor

    return lineas, tabla_valores

def extraer_valor(lineas, tabla_valores, etiqueta):
    etiqueta_norm = normalizar_texto(etiqueta)

    for key in tabla_valores:
        if etiqueta_norm in key or key in etiqueta_norm:
            return tabla_valores[key]

    for linea in lineas:
        linea_norm = normalizar_texto(linea)
        if etiqueta_norm in linea_norm or linea_norm in etiqueta_norm:
            if ':' in linea:
                return linea.split(':', 1)[1].strip()
            return linea.replace(etiqueta, '').strip()
    
    return ""

def extraer_negocios_hijos(lineas):
    hijos = []
    for linea in lineas:
        if "negocio hijo" in linea.lower():
            partes = linea.split(":")
            if len(partes) >= 2:
                nombre = partes[1].strip()
                if nombre:
                    hijos.append({"nombre": nombre})
    return hijos

def procesar_documentos(carpeta):
    resultados = []
    for archivo in os.listdir(carpeta):
        if archivo.endswith(".docx"):
            ruta_archivo = os.path.join(carpeta, archivo)
            print(f"Procesando archivo: {archivo}")
            doc = Document(ruta_archivo)
            lineas, tabla_valores = extraer_lineas_y_tablas(doc)

            hijos = extraer_negocios_hijos(lineas)

            data = {
                "nombre": extraer_valor(lineas, tabla_valores, "Nombre del negocio"),
                "direccion": extraer_valor(lineas, tabla_valores, "Dirección de la casa matriz"),
                "propietario": extraer_valor(lineas, tabla_valores, "Nombre del propietario"),
                "admin": "",
                "tel_propietario": extraer_valor(lineas, tabla_valores, "Teléfono del propietario"),
                "tel_admin": "",
                "comercial": extraer_valor(lineas, tabla_valores, "Comercial responsable"),
                "licencia": extraer_valor(lineas, tabla_valores, "Mensualidad contratada"),
                "tipo": "casa_matriz",
                "negocios_hijos": json.dumps(hijos) if hijos else ""
            }

            resultados.append(data)
    return resultados

# Procesamiento y envío a la API
carpeta = "ficha bases de clientes"
resultados = procesar_documentos(carpeta)

for data in resultados:
    print("=== Datos extraídos ===")
    print(data)

    url = "http://127.0.0.1:5000/api/agregar_negocio"
    response = requests.post(url, json=data)
    print("Código respuesta:", response.status_code)
    try:
        print("Respuesta JSON:", response.json())
    except Exception:
        print("Respuesta texto:", response.text)
